from fastapi import FastAPI, Depends, File, UploadFile, HTTPException, Request, BackgroundTasks, status # type: ignore
from fastapi.middleware.cors import CORSMiddleware # type: ignore
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials # type: ignore
from pydantic import BaseModel, EmailStr
import typing as t
import uvicorn # type: ignore
import os
import fitz  # PyMuPDF
import logging
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
from fastapi import FastAPI, HTTPException, Depends, Request, Response, Cookie, status, File, UploadFile # type: ignore
import asyncio
from concurrent.futures import ProcessPoolExecutor, ThreadPoolExecutor
import multiprocessing as mp
from functools import partial
import time
import sqlite3
from datetime import datetime, timedelta
from passlib.context import CryptContext
from contextlib import asynccontextmanager
import tempfile
import shutil
import gc
import re
import random
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import io


from qdrant_engine import QdrantIndex
from config import settings
import os
import sqlite3
import logging
import asyncio
from datetime import datetime, timedelta
from typing import Optional
import hashlib
import secrets

from fastapi import FastAPI, HTTPException, Depends, Request, Response, status
from fastapi.responses import JSONResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, EmailStr
from passlib.context import CryptContext
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Logging Configuration
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# Session Configuration
SECRET_KEY = getattr(settings, '772jbCe6yK1dqYViLbp5twgUkAIXncCR', "00UlaD9ESbxHs-QAZvhOX3iINp-SHazLSe4FEuf0BKieTsrfTIISASyYaBYmFS5V")
# Session configuration
SESSION_EXPIRE_MINUTES = 1440  # 24 hours
COOKIE_NAME = "session_token"

# Password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# Database Configuration
DATABASE_PATH = "app/database/users.db"

# Pydantic Models
class UserCreate(BaseModel):
    username: str
    email: EmailStr
    password: str
    full_name: str = None

class UserLogin(BaseModel):
    username: str
    password: str

class UserResponse(BaseModel):
    id: str
    username: str
    email: str
    full_name: str = None
    is_active: bool
    created_at: datetime

class LoginResponse(BaseModel):
    message: str
    user: UserResponse

class UserQuery(BaseModel):
    query: str

class ComprehensiveQuery(BaseModel):
    query: str
    use_pdf_context: bool = True
    generate_comprehensive: bool = True
    max_tokens: int = 2000

def check_and_migrate_database():
    """Check if database schema needs migration and update if necessary"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    
    try:
        cursor.execute("PRAGMA table_info(user_sessions)")
        columns = cursor.fetchall()
        column_names = [column[1] for column in columns]

        needs_commit = False

        if 'session_token' not in column_names:
            logging.info("session_token column missing, adding it...")
            cursor.execute('''
                ALTER TABLE user_sessions 
                ADD COLUMN session_token TEXT UNIQUE
            ''')
            needs_commit = True
        
        if 'expires_at' not in column_names:
            logging.info("expires_at column missing, adding it...")
            cursor.execute('''
                ALTER TABLE user_sessions 
                ADD COLUMN expires_at TIMESTAMP
            ''')
            needs_commit = True

        if 'current_file' not in column_names:
            logging.info("current_file column missing, adding it...")
            cursor.execute('''
                ALTER TABLE user_sessions 
                ADD COLUMN current_file TEXT
            ''')
            needs_commit = True

        if needs_commit:
            conn.commit()
            logging.info("Database schema updated successfully")
        else:
            logging.info("Database schema is up to date")
            
    except sqlite3.OperationalError as e:
        logging.error(f"Database migration error: {e}")
        recreate_sessions_table(cursor)
        conn.commit()
    finally:
        conn.close()

def recreate_sessions_table(cursor):
    """Recreate the user_sessions table with correct schema"""
    cursor.execute('DROP TABLE IF EXISTS user_sessions')
    cursor.execute('''
        CREATE TABLE user_sessions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            session_token TEXT NOT NULL UNIQUE,
            expires_at TIMESTAMP NOT NULL,
            current_file TEXT,  -- NEW FIELD
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id) ON DELETE CASCADE
        )
    ''')
    logging.info("user_sessions table recreated with correct schema (including current_file)")

def init_database():
    """Initialize SQLite database and create tables"""
    os.makedirs(os.path.dirname(DATABASE_PATH), exist_ok=True)
    
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    
    # Create users table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            email TEXT UNIQUE NOT NULL,
            full_name TEXT,
            hashed_password TEXT NOT NULL,
            is_active BOOLEAN DEFAULT TRUE,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Create sessions table for cookie session management
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS user_sessions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            session_token TEXT NOT NULL UNIQUE,
            expires_at TIMESTAMP NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id) ON DELETE CASCADE
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS conversations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            query TEXT NOT NULL,
            response TEXT NOT NULL,
            timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id) ON DELETE CASCADE
        )
    ''')
    
    conn.commit()
    conn.close()
    logging.info("Database initialized successfully")
    
    # Run migration check after initialization
    check_and_migrate_database()
    
def update_current_file_for_session(session_token: str, filename: str):
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        token_hash = hash_session_token(session_token)
        cursor.execute('''
            UPDATE user_sessions SET current_file = ?
            WHERE session_token = ?
        ''', (filename, token_hash))
        conn.commit()
    finally:
        conn.close()

def get_current_file_for_session(session_token: str) -> Optional[str]:
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        token_hash = hash_session_token(session_token)
        cursor.execute('''
            SELECT current_file FROM user_sessions
            WHERE session_token = ?
        ''', (token_hash,))
        result = cursor.fetchone()
        return result[0] if result and result[0] else None
    finally:
        conn.close()

    
def classify_question_type(query: str) -> str:
    """Classify the type of question being asked."""
    query_lower = query.lower()
    if any(word in query_lower for word in ['what is', 'what are', 'define']):
        return "definition"
    elif any(word in query_lower for word in ['how to', 'how do', 'how can']):
        return "procedure"
    elif any(word in query_lower for word in ['why', 'because', 'reason']):
        return "explanation"
    elif any(word in query_lower for word in ['when', 'date', 'time']):
        return "temporal"
    elif any(word in query_lower for word in ['where', 'location']):
        return "location"
    else:
        return "general"


def create_user_session(user_id: int) -> str:
    """Create a new session for the user"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    
    try:
        # Generate a secure random session token
        session_token = secrets.token_urlsafe(32)
        
        # Calculate expiration time (assuming SESSION_EXPIRE_MINUTES is defined)
        SESSION_EXPIRE_MINUTES = 1440  # 24 hours, adjust as needed
        expires_at = datetime.now() + timedelta(minutes=SESSION_EXPIRE_MINUTES)
        
        # Insert new session
        cursor.execute('''
            INSERT INTO user_sessions (user_id, session_token, expires_at)
            VALUES (?, ?, ?)
        ''', (user_id, session_token, expires_at))
        
        conn.commit()
        logging.info(f"Created session for user {user_id}")
        return session_token
        
    except sqlite3.IntegrityError as e:
        logging.error(f"Session creation error: {e}")
        # If token collision (very unlikely), generate a new one
        conn.close()
        return create_user_session(user_id)
    except Exception as e:
        logging.error(f"Unexpected error creating session: {e}")
        raise
    finally:
        conn.close()

def get_user_from_session(session_token: str):
    """Get user data from session token"""
    print("session token is here ",  session_token)
    conn = sqlite3.connect(DATABASE_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    try:
        cursor.execute('''
            SELECT u.id, u.username, u.email, u.full_name, u.is_active, u.created_at
            FROM users u
            JOIN user_sessions s ON u.id = s.user_id
            WHERE s.session_token = ? 
            AND s.expires_at > CURRENT_TIMESTAMP
        ''', (session_token,))
        
        user = cursor.fetchone()
        print("USer is here ", user)
        if user:
            return dict(user)
        return None
        
    except Exception as e:
        logging.error(f"Error getting user from session: {e}")
        return None
    finally:
        conn.close()

def cleanup_expired_sessions():
    """Remove expired sessions"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    
    try:
        cursor.execute('''
            DELETE FROM user_sessions 
            WHERE expires_at < CURRENT_TIMESTAMP
        ''')
        
        deleted_count = cursor.rowcount
        conn.commit()
        
        if deleted_count > 0:
            logging.info(f"Cleaned up {deleted_count} expired sessions")
            
    except Exception as e:
        logging.error(f"Error cleaning up sessions: {e}")
    finally:
        conn.close()

def delete_user_session(session_token: str):
    """Delete a specific session (for logout)"""
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    
    try:
        cursor.execute('''
            DELETE FROM user_sessions 
            WHERE session_token = ?
        ''', (session_token,))
        
        conn.commit()
        logging.info("Session deleted successfully")
        
    except Exception as e:
        logging.error(f"Error deleting session: {e}")
    finally:
        conn.close()

# Add this to your main.py startup
def startup_database():
    """Initialize database and run migrations on startup"""
    init_database()
    cleanup_expired_sessions()  # Clean up any expired sessions on startup

def get_db_connection():
    """Get database connection"""
    conn = sqlite3.connect(DATABASE_PATH)
    conn.row_factory = sqlite3.Row
    return conn

# Password utilities
def verify_password(plain_password: str, hashed_password: str) -> bool:
    """Verify password against hash"""
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password: str) -> str:
    """Hash password"""
    return pwd_context.hash(password)

# Session utilities
def create_session_token() -> str:
    """Create a secure session token"""
    return secrets.token_urlsafe(32)

def hash_session_token(token: str) -> str:
    """Hash session token for storage"""
    return hashlib.sha256((token + SECRET_KEY).encode()).hexdigest()

def verify_session_token(token: str, stored_hash: str) -> bool:
    """Verify session token against stored hash"""
    return hash_session_token(token) == stored_hash

# Database operations
def create_user(user_data: UserCreate) -> dict:
    """Create a new user in database"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        # Check if user already exists
        cursor.execute("SELECT id FROM users WHERE username = ? OR email = ?", 
                      (user_data.username, user_data.email))
        if cursor.fetchone():
            raise HTTPException(
                status_code=400,
                detail="Username or email already registered"
            )
        
        # Hash password and create user
        hashed_password = get_password_hash(user_data.password)
        cursor.execute('''
            INSERT INTO users (username, email, full_name, hashed_password)
            VALUES (?, ?, ?, ?)
        ''', (user_data.username, user_data.email, user_data.full_name, hashed_password))
        
        user_id = cursor.lastrowid
        conn.commit()
        
        # Fetch created user
        cursor.execute("SELECT * FROM users WHERE id = ?", (user_id,))
        user = cursor.fetchone()
        
        return dict(user)
        
    except sqlite3.IntegrityError:
        raise HTTPException(
            status_code=400,
            detail="Username or email already registered"
        )
    finally:
        conn.close()

def authenticate_user(username: str, password: str):
    """Authenticate user credentials"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        cursor.execute("SELECT * FROM users WHERE username = ?", (username,))
        user = cursor.fetchone()
        
        if not user:
            return False
        
        if not verify_password(password, user['hashed_password']):
            return False
        
        return dict(user)
        
    finally:
        conn.close()

def get_user_by_username(username: str):
    """Get user by username"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        cursor.execute("SELECT * FROM users WHERE username = ? AND is_active = TRUE", (username,))
        user = cursor.fetchone()
        return dict(user) if user else None
    finally:
        conn.close()

def create_user_session(user_id: int) -> str:
    """Create a new session for user"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        # Generate session token
        session_token = create_session_token()
        token_hash = hash_session_token(session_token)
        expires_at = datetime.utcnow() + timedelta(minutes=SESSION_EXPIRE_MINUTES)
        
        # Store session in database
        cursor.execute('''
            INSERT INTO user_sessions (user_id, session_token, expires_at)
            VALUES (?, ?, ?)
        ''', (user_id, token_hash, expires_at))
        conn.commit()
        
        return session_token
        
    finally:
        conn.close()

def get_user_by_session(session_token: str):
    """Get user by session token"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        token_hash = hash_session_token(session_token)
        
        cursor.execute('''
            SELECT u.*, s.expires_at
            FROM users u
            JOIN user_sessions s ON u.id = s.user_id
            WHERE s.session_token = ? AND s.expires_at > ? AND u.is_active = TRUE
        ''', (token_hash, datetime.utcnow()))
        
        result = cursor.fetchone()
        return dict(result) if result else None
        
    finally:
        conn.close()

def invalidate_session(session_token: str):
    """Invalidate a session"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        token_hash = hash_session_token(session_token)
        cursor.execute("DELETE FROM user_sessions WHERE session_token = ?", (token_hash,))
        conn.commit()
        return cursor.rowcount > 0
    finally:
        conn.close()

def invalidate_all_user_sessions(user_id: int):
    """Invalidate all sessions for a user"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        cursor.execute("DELETE FROM user_sessions WHERE user_id = ?", (user_id,))
        conn.commit()
        return cursor.rowcount
    finally:
        conn.close()

# Add this function to store conversations
def store_conversation(user_id: int, query: str, response: str):
    """Store conversation in database"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        cursor.execute('''
            INSERT INTO conversations (user_id, query, response)
            VALUES (?, ?, ?)
        ''', (user_id, query, response))
        conn.commit()
    except Exception as e:
        logging.error(f"Error storing conversation: {str(e)}")
    finally:
        conn.close()
        
# Add this function to get conversation history
def get_conversation_history(user_id: int = None, limit: int = 100):
    """Get conversation history from database"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        if user_id:
            cursor.execute('''
                SELECT query, response, timestamp FROM conversations 
                WHERE user_id = ? 
                ORDER BY timestamp DESC 
                LIMIT ?
            ''', (user_id, limit))
        else:
            cursor.execute('''
                SELECT query, response, timestamp FROM conversations 
                ORDER BY timestamp DESC 
                LIMIT ?
            ''', (limit,))
        
        conversations = cursor.fetchall()
        return [dict(conv) for conv in conversations]
    except Exception as e:
        logging.error(f"Error getting conversation history: {str(e)}")
        return []
    finally:
        conn.close()

# Function to create Word document with conversations
def create_conversation_document(conversations: list, user_name: str = "User"):
    """Create a Word document with conversation history"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Conversation History', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"User: {user_name}")
    doc.add_paragraph(f"Total Conversations: {len(conversations)}")
    doc.add_paragraph("=" * 50)
    
    # Add each conversation
    for i, conv in enumerate(conversations, 1):
        # Add conversation number
        conv_heading = doc.add_heading(f'Conversation {i}', level=1)
        
        # Add timestamp
        timestamp = conv.get('timestamp', 'Unknown')
        if isinstance(timestamp, str):
            try:
                timestamp = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                timestamp = timestamp.strftime('%Y-%m-%d %H:%M:%S')
            except:
                pass
        
        doc.add_paragraph(f"Time: {timestamp}")
        
        # Add query
        query_para = doc.add_paragraph()
        query_run = query_para.add_run("Query: ")
        query_run.bold = True
        cleaned_query = clean_text_for_docx(conv.get('query', 'No query available'))
        query_para.add_run(cleaned_query)
        # query_para.add_run(conv.get('query', 'No query available'))
        
        # Add response
        response_para = doc.add_paragraph()
        response_run = response_para.add_run("Response: ")
        response_run.bold = True
        cleaned_response = clean_text_for_docx(conv.get('response', 'No response available'))
        response_para.add_run(cleaned_response)
        
        # Add separator
        doc.add_paragraph("-" * 50)
    
    return doc

def cleanup_expired_sessions():
    """Clean up expired sessions"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        cursor.execute("DELETE FROM user_sessions WHERE expires_at < ?", (datetime.utcnow(),))
        conn.commit()
        logging.info(f"Cleaned up {cursor.rowcount} expired sessions")
    finally:
        conn.close()

# Dependency for protected routes
async def get_current_user(session_token: str = Cookie(None, alias=COOKIE_NAME)):
    """Get current user from session token"""
    
    if not session_token:
        print("no session token found")
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="No session token found"
        )
    
    user = get_user_by_session(session_token)
    if not user:
        print("no user with session found")
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid or expired session"
        )
    
    return user
# Optional dependency for routes that can work with or without auth
async def get_current_user_optional(request: Request):
    """Get current user if authenticated, None otherwise"""
    try:
        return await get_current_user(request)
    except HTTPException:
        return None

# Background task for periodic cleanup
async def periodic_cleanup():
    """Periodic cleanup of expired sessions"""
    while True:
        try:
            cleanup_expired_sessions()
            # Sleep for 1 hour
            await asyncio.sleep(3600)
        except Exception as e:
            logging.error(f"Error in periodic cleanup: {str(e)}")
            await asyncio.sleep(3600)

async def start_periodic_cleanup():
    """Start the periodic cleanup as a background task"""
    cleanup_task = asyncio.create_task(periodic_cleanup())
    return cleanup_task

@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    init_database()
    cleanup_expired_sessions()
    cleanup_old_files()
    
    # Start the periodic cleanup task
    cleanup_task = asyncio.create_task(periodic_cleanup())
    
    yield
    
    # Shutdown - cancel the cleanup task
    cleanup_task.cancel()
    try:
        await cleanup_task
    except asyncio.CancelledError:
        pass
            
app = FastAPI(
    title="DrQA Backend API with Authentication",
    docs_url="/docs",
    lifespan=lifespan
)

# CORS settings
origins = [
    "http://localhost:8000",
    "http://localhost:3000",
    "http://127.0.0.1:8000",
    "http://127.0.0.1:3000",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # ✅ replace with your frontend's URL
    allow_credentials=True,                   # ✅ this enables cookies
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load the Qdrant index
qdrant_index = QdrantIndex(settings.qdrant_host, settings.qdrant_api_key, False)

# Ensure the documents directory exists
os.makedirs(os.path.join("app", "documents"), exist_ok=True)

# Global process pool for CPU-intensive tasks
process_pool = ProcessPoolExecutor(max_workers=mp.cpu_count())

@app.post("/auth/register", response_model=UserResponse)
async def register(user_data: UserCreate):
    """Register a new user"""
    try:
        user = create_user(user_data)
        return UserResponse(
            id=str(user['id']),
            username=user['username'],
            email=user['email'],
            full_name=user['full_name'],
            is_active=user['is_active'],
            created_at=datetime.fromisoformat(user['created_at'])
        )
    except Exception as e:
        if "already registered" in str(e):
            raise e
        logging.error(f"Registration error: {str(e)}")
        raise HTTPException(status_code=500, detail="Registration failed")

@app.on_event("startup")
async def startup_event():
    """Initialize database and clean up expired sessions on startup"""
    startup_database()
    logging.info("Application started successfully")


# Updated download-file route
@app.get("/download-conversations")
async def download_conversations(
    user_id: int = None,
    limit: int = 100,
    current_user: dict = Depends(get_current_user)
):
    """Download conversation history as Word document"""
    try:
        # Get conversation history
        conversations = get_conversation_history(user_id, limit)
        
        if not conversations:
            raise HTTPException(status_code=404, detail="No conversations found")
        
        # Determine user name
        user_name = "Anonymous"
        if current_user:
            user_name = current_user.get('username', 'User')
        
        # Create Word document
        doc = create_conversation_document(conversations, user_name)
        
        # Save to memory buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"conversation_history_{timestamp}.docx"
        
        # Create streaming response
        def generate():
            yield buffer.read()
        
        return StreamingResponse(
            io.BytesIO(buffer.read()),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        logging.error(f"Error downloading conversations: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")



# Updated download-file route
@app.get("/download-conversations")
async def download_conversations(
    request: Request,
    user_id: int = None,
    limit: int = 100
):
    """Download conversation history as Word document"""
    try:
        # Get current user if authenticated
        current_user = await get_current_user(request)
        
        # If user_id is not provided and user is authenticated, use current user's ID
        if user_id is None and current_user:
            user_id = current_user['id']
        
        # Get conversation history
        conversations = get_conversation_history(user_id, limit)
        
        if not conversations:
            raise HTTPException(status_code=404, detail="No conversations found")
        
        # Determine user name
        user_name = "Anonymous"
        if current_user:
            user_name = current_user.get('username', 'User')
        
        # Create Word document
        doc = create_conversation_document(conversations, user_name)
        
        # Save to memory buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"conversation_history_{timestamp}.docx"
        
        # Return streaming response
        return StreamingResponse(
            io.BytesIO(buffer.read()),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        logging.error(f"Error downloading conversations: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")



# Updated login endpoint
@app.post("/auth/login", response_model=LoginResponse)
async def login(user_credentials: UserLogin, response: Response):
    """Authenticate user and set session cookie"""
    try:
        user = authenticate_user(user_credentials.username, user_credentials.password)
        
        if not user:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Incorrect username or password"
            )
        
        # Create session
        print('user_id  ', user['id'])
        session_token = create_user_session(user['id'])
        
        # Set cookie
        response.set_cookie(
            key=COOKIE_NAME,
            value=session_token,
            max_age=SESSION_EXPIRE_MINUTES * 60,  # Convert to seconds
            httponly=True,  # Prevent XSS attacks
            secure=False,   # Set to True for HTTPS in production
            samesite="lax"  # CSRF protection
        )
        
        return LoginResponse(
            message="Login successful",
            user=UserResponse(
                id=str(user['id']),
                username=user['username'],
                email=user['email'],
                full_name=user['full_name'],
                is_active=user['is_active'],
                created_at=datetime.fromisoformat(user['created_at'])
            )
        )
        
    except Exception as e:
        logging.error(f"Login error: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Internal server error"
        )
        
        
@app.get("/auth/profile", response_model=UserResponse)
async def get_profile(current_user: dict = Depends(get_current_user)):
    """Get current user profile"""
    return UserResponse(
        id=str(current_user['id']),
        username=current_user['username'],
        email=current_user['email'],
        full_name=current_user['full_name'],
        is_active=current_user['is_active'],
        created_at=datetime.fromisoformat(current_user['created_at'])
    )
    
        
async def logout(request: Request, response: Response):
    """Logout user and invalidate session"""
    session_token = request.cookies.get(COOKIE_NAME)
    
    if session_token:
        invalidate_session(session_token)
    
    # Clear cookie
    response.delete_cookie(key=COOKIE_NAME)
    
    return {"message": "Successfully logged out"}

# Public Routes (no authentication required)
@app.get("/")
async def root():
    return {"message": "Server is up and running!"}

@app.get("/test-upload")
async def test_upload():
    """Test endpoint to verify API is working"""
    return {"message": "Upload endpoint is accessible", "status": "ok"}

def update_current_file_for_session(session_token: str, filename: str):
    """Update the current working PDF file for the session."""
    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        token_hash = hash_session_token(session_token)
        cursor.execute('''
            UPDATE user_sessions
            SET current_file = ?
            WHERE session_token = ?
        ''', (filename, token_hash))
        conn.commit()
        logging.info(f"Session updated with current_file = {filename}")
    except Exception as e:
        logging.error(f"Failed to update current file for session: {str(e)}")
    finally:
        conn.close()


# File Upload Route
@app.post("/upload-file")
async def upload_file(request: Request, file: UploadFile = File(...)):
    """Upload PDF file, validate it, index it, and set it as current working file in session."""
    try:
        logging.info(f"Received file upload request: {file.filename}")
        
        if not file or not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        if not file.filename.lower().endswith(".pdf"):
            raise HTTPException(status_code=400, detail="Only PDF files are allowed")

        # Prepare save path
        documents_dir = "app/documents"
        os.makedirs(documents_dir, exist_ok=True)

        file_path = os.path.join(documents_dir, file.filename)
        file_path = os.path.abspath(file_path)

        # Rename if exists
        if os.path.exists(file_path):
            timestamp = int(time.time())
            name, ext = os.path.splitext(file.filename)
            new_filename = f"{name}_{timestamp}{ext}"
            file_path = os.path.join(documents_dir, new_filename)
            logging.info(f"File exists; renamed to: {new_filename}")
        else:
            new_filename = file.filename

        # Save file
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)

        # Validate metadata (optional)
        if 'validate_and_fix_pdf_metadata' in globals():
            validate_and_fix_pdf_metadata(file_path)

        # Index with Qdrant
        if 'qdrant_index' in globals():
            await qdrant_index.insert_into_index_async(file_path, new_filename)
        else:
            raise HTTPException(status_code=500, detail="Qdrant index not initialized")

        # Set this file as current_file for the user session
        session_token = request.cookies.get(COOKIE_NAME)
        if not session_token:
            raise HTTPException(status_code=401, detail="No session token found. Please login first.")
        update_current_file_for_session(session_token, new_filename)

        return {
            "message": f"File '{new_filename}' uploaded, indexed, and set as active.",
            "filename": new_filename,
            "status": "success"
        }

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Unexpected error in file upload: {str(e)}")
        raise HTTPException(status_code=500, detail=f"File upload failed: {str(e)}")

# Add a simple test endpoint to verify the API is working
@app.get("/test-upload")
async def test_upload():
    """Test endpoint to verify API is working"""
    return {"message": "Upload endpoint is accessible", "status": "ok"}

# Add endpoint to check upload requirements
@app.get("/upload-info")
async def upload_info():
    """Get information about upload requirements"""
    return {
        "accepted_types": ["application/pdf"],
        "max_file_size": "No limit specified",
        "required_fields": ["file"],
        "endpoint": "/upload-file"
    }

@app.post("/comprehensive-query")
async def comprehensive_query(
    input_query: ComprehensiveQuery,
    request: Request,
    current_user: dict = Depends(get_current_user)
):
    """
    Enhanced query endpoint that provides comprehensive responses,
    stores conversation, and uses the current working PDF file for context.
    """
    try:
        start_time = time.time()

        # Step 1: Retrieve previous conversations
        conversations = get_conversation_history(user_id=current_user["id"], limit=5)
        conversation_history = "\n".join([f"Q: {conv['query']}\nA: {conv['response']}" for conv in conversations])

        # Step 2: Determine the current working file from session
        session_token = request.cookies.get(COOKIE_NAME)
        current_file = get_current_file_for_session(session_token) if session_token else None

        if not current_file:
            raise HTTPException(status_code=400, detail="No active file found. Please upload or select a file.")

        # Step 3: Load PDF context from the current file (if enabled)
        pdf_context = ""
        pdf_sources = []

        if input_query.use_pdf_context:
            try:
                full_context = qdrant_index._get_specific_pdf_content(current_file)
                pdf_context = full_context[:12000]  # truncate if necessary
                logging.info(f"Loaded context from current file '{current_file}' ({len(pdf_context)} characters)")
            except Exception as e:
                logging.warning(f"Could not load context from current file: {str(e)}")

        # Step 4: Combine all context elements for AI prompt
        full_prompt_context = f"Previous Conversations:\n{conversation_history}\n\nCurrent Query: {input_query.query}\n"
        if pdf_context:
            full_prompt_context += f"\nPDF Context from {current_file}:\n{pdf_context}"

        # Step 5: Generate the comprehensive AI response
        comprehensive_response = await generate_comprehensive_response(
            query=input_query.query,
            pdf_context=full_prompt_context,
            max_tokens=input_query.max_tokens
        )

        # Step 6: Store query and response
        store_conversation(current_user["id"], input_query.query, comprehensive_response)

        # Step 7: Build the response payload
        final_response = {
            "query": input_query.query,
            "comprehensive_answer": comprehensive_response,
            "pdf_filename": current_file,
            "pdf_context_used": bool(pdf_context),
            "response_time": time.time() - start_time,
            "timestamp": time.time(),
        }

        # Optional: Extract insights
        if pdf_context:
            final_response["combined_insights"] = await combine_pdf_and_ai_insights(
                input_query.query, pdf_context, comprehensive_response
            )

        return final_response

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Error in comprehensive query: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Query processing failed: {str(e)}")

@app.post("/query")
async def query_index(input_query: UserQuery, request: Request):
    """
    Standard query endpoint with enhanced response formatting - Protected Route
    Requires valid session cookie for authentication
    """
    try:
        # Check for authentication cookie
        session_token = request.cookies.get(COOKIE_NAME)
        if not session_token:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Authentication required. Please login."
            )
        
        # Validate session token
        user = get_user_by_session(session_token)
        if not user:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Invalid or expired session. Please login again."
            )
        
        # Log the authenticated query
        logging.info(f"Received query from user {user['username']} (ID: {user['id']}): {input_query.query}")
        
        # If query_and_generate_response is async, await it
        # If it's sync, you might want to run it in a thread pool for better performance
        try:
            # Try async first
            result = await qdrant_index.query_and_generate_response(query=input_query.query)
        except TypeError:
            # If it's not async, run in thread pool to avoid blocking
            import asyncio
            from concurrent.futures import ThreadPoolExecutor
            
            loop = asyncio.get_event_loop()
            with ThreadPoolExecutor() as executor:
                result = await loop.run_in_executor(
                    executor, 
                    qdrant_index.query_and_generate_response, 
                    input_query.query
                )
        
        if isinstance(result, tuple) and len(result) >= 2:
            generated_response, relevant_docs = result[0], result[1]
        else:
            generated_response = str(result)
            relevant_docs = []
        
        # Clean and format response
        cleaned_response = format_for_frontend(generated_response)
        
        relevant_docs_content = []
        if relevant_docs:
            relevant_docs = [doc for doc in relevant_docs if hasattr(doc, "page_content") and doc.page_content]
            relevant_docs_content = [simple_clean_text(doc.page_content) for doc in relevant_docs]
        
        return {
            "response": cleaned_response["main_answer"],
            "timestamp": datetime.utcnow().isoformat(),
            "user_id": str(user['id'])  # Optional: include user context
        }
    
    except HTTPException:
        # Re-raise HTTP exceptions (authentication errors)
        raise
    except Exception as e:
        logging.error(f"Query processing error for user {user.get('username', 'unknown') if 'user' in locals() else 'unauthenticated'}: {str(e)}")
        raise HTTPException(status_code=500, detail="Error processing the query.")


@app.post("/auth/login", response_model=LoginResponse)
async def login(user_credentials: UserLogin, response: Response):
    """Authenticate user and set session cookie"""
    user = authenticate_user(user_credentials.username, user_credentials.password)
    
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password"
        )
    
    # Create session
    session_token = create_user_session(user['id'])
    
    # Set cookie
    response.set_cookie(
        key=COOKIE_NAME,
        value=session_token,
        max_age=SESSION_EXPIRE_MINUTES * 60,  # Convert to seconds
        httponly=True,  # Prevent XSS attacks
        secure=False,   # Set to True for HTTPS in production
        samesite="lax"  # CSRF protection
    )
    
    return LoginResponse(
        message="Login successful",
        user=UserResponse(
            id=str(user['id']),
            username=user['username'],
            email=user['email'],
            full_name=user['full_name'],
            is_active=user['is_active'],
            created_at=datetime.fromisoformat(user['created_at'])
        )
    )

@app.post("/auth/logout")
async def logout(response: Response, session_token: str = Cookie(None, alias=COOKIE_NAME)):
    """Logout user and clear session"""
    if session_token:
        delete_user_session(session_token)
    
    response.delete_cookie(key=COOKIE_NAME)
    return {"message": "Logged out successfully"}

# Alternative version if you know query_and_generate_response is sync
@app.post("/query-sync")
async def query_index_sync(input_query: UserQuery):
    """
    Query endpoint optimized for sync query_and_generate_response method
    """
    try:
        logging.info(f"Received query: {input_query.query}")
        
        # Run sync method in thread pool to avoid blocking the event loop
        import asyncio
        from concurrent.futures import ThreadPoolExecutor
        
        loop = asyncio.get_event_loop()
        with ThreadPoolExecutor() as executor:
            result = await loop.run_in_executor(
                executor, 
                qdrant_index.query_and_generate_response, 
                input_query.query
            )
        
        if isinstance(result, tuple) and len(result) >= 2:
            generated_response, relevant_docs = result[0], result[1]
        else:
            generated_response = str(result)
            relevant_docs = []
        
        # Clean and format response
        cleaned_response = format_for_frontend(generated_response)
        
        relevant_docs_content = []
        if relevant_docs:
            relevant_docs = [doc for doc in relevant_docs if hasattr(doc, "page_content") and doc.page_content]
            relevant_docs_content = [simple_clean_text(doc.page_content) for doc in relevant_docs]
        
        return {
            "response": cleaned_response["main_answer"],
            "formatted_response": cleaned_response,
            "relevant_docs": relevant_docs_content,
            "total_sources": len(relevant_docs_content),
            "timestamp": datetime.utcnow().isoformat()
        }

    except Exception as e:
        logging.error(f"Query processing error: {str(e)}")
        raise HTTPException(status_code=500, detail="Error processing the query.")


# If you want to make query_and_generate_response async as well
async def query_and_generate_response_async(self, query: str):
    """
    Async version of query_and_generate_response method
    """
    try:
        # Run potentially blocking operations in thread pool
        import asyncio
        from concurrent.futures import ThreadPoolExecutor
        
        loop = asyncio.get_event_loop()
        
        # If your retrieval/generation involves blocking operations, wrap them
        with ThreadPoolExecutor() as executor:
            # Example: if you have blocking retrieval
            # relevant_docs = await loop.run_in_executor(executor, self.retrieve_documents, query)
            
            # Or if the entire method is blocking:
            result = await loop.run_in_executor(
                executor, 
                self._sync_query_and_generate_response, 
                query
            )
            
        return result
        
    except Exception as e:
        logging.error(f"Async query processing error: {str(e)}")
        raise

# File management endpoints (Protected)
@app.delete("/remove-file/{filename}")
async def remove_file(filename: str, request: Request):
    try:
        # Check for authentication cookie
        session_token = request.cookies.get(COOKIE_NAME)
        if not session_token:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Authentication required. Please login."
            )
        
        # Validate session token
        user = get_user_by_session(session_token)
        if not user:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Invalid or expired session. Please login again."
            )
        
        file_path = os.path.join("app", "documents", os.path.basename(filename))
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found")
        
        qdrant_index.remove_document(filename)
        os.remove(file_path)
        
        return {
            "status": "success", 
            "message": f"File {filename} removed successfully",
        }
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Error removing file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error removing file: {str(e)}")

@app.get("/download-file/{filename}")
async def download_file(filename: str, request: Request):
    # Check for authentication cookie
    session_token = request.cookies.get(COOKIE_NAME)
    if not session_token:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Authentication required. Please login."
        )
    
    # Validate session token
    user = get_user_by_session(session_token)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid or expired session. Please login again."
        )
    
    file_path = os.path.join("app", "documents", os.path.basename(filename))
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        file_path,
        media_type="application/pdf",
        filename=filename
    )

@app.get("/list-files")
async def list_files(request: Request):
    try:
        # Check for authentication cookie
        session_token = request.cookies.get(COOKIE_NAME)
        if not session_token:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Authentication required. Please login."
            )
        
        # Validate session token
        user = get_user_by_session(session_token)
        if not user:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Invalid or expired session. Please login again."
            )
        
        documents_dir = os.path.join("app", "documents")
        files = [f for f in os.listdir(documents_dir) if f.endswith(".pdf")]
        return {"files": files}
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Error listing files: {str(e)}")
        raise HTTPException(status_code=500, detail="Error listing files")
    
    
# Utility functions (keep all existing ones)
def validate_and_fix_pdf_metadata(pdf_path):
    """
    Validate and fix PDF metadata to ensure all standard fields are present.
    """
    try:
        doc = fitz.open(pdf_path)
        metadata = doc.metadata

        required_fields = {
            'title': 'Default Title',
            'author': 'Unknown Author',
            'subject': 'Uploaded Document',
            'keywords': 'document,upload',
            'creator': 'DrQA System',
            'producer': 'DrQA PDF Processor',
        }

        for field, default in required_fields.items():
            if not metadata.get(field):
                metadata[field] = default

        if not metadata.get('creationDate'):
            metadata['creationDate'] = fitz.get_pdf_now()
        if not metadata.get('modDate'):
            metadata['modDate'] = fitz.get_pdf_now()

        doc.set_metadata(metadata)

        temp_path = f"{pdf_path}_fixed.pdf"
        doc.save(temp_path, 
                garbage=4,
                deflate=True,
                no_new_id=True
        )
        doc.close()

        os.replace(temp_path, pdf_path)

        with fitz.open(pdf_path) as verify_doc:
            logging.info(f"Verified PDF metadata: {verify_doc.metadata}")

        logging.info(f"Successfully processed PDF: {pdf_path}")

    except Exception as e:
        logging.error(f"PDF processing failed: {str(e)}")
        raise HTTPException(status_code=500, detail="PDF metadata validation failed.")

def process_pdf_chunk(args):
    """
    Process a chunk of PDF pages in a separate process.
    This function will be executed in parallel.
    """
    chunk_pages, pdf_path, start_page, chunk_id = args
    
    try:
        import fitz
        from langchain_community.document_loaders import PDFMinerLoader
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        
        # Create a temporary PDF with just this chunk
        temp_pdf_path = f"{pdf_path}_chunk_{chunk_id}.pdf"
        
        # Open original PDF and create chunk
        original_doc = fitz.open(pdf_path)
        chunk_doc = fitz.open()
        
        for page_num in range(start_page, min(start_page + chunk_pages, original_doc.page_count)):
            chunk_doc.insert_pdf(original_doc, from_page=page_num, to_page=page_num)
        
        chunk_doc.save(temp_pdf_path)
        chunk_doc.close()
        original_doc.close()
        
        # Process the chunk
        loader = PDFMinerLoader(temp_pdf_path)
        docs = loader.load()
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        documents = text_splitter.split_documents(docs)
        
        # Clean up temporary file
        os.remove(temp_pdf_path)
        
        # Return processed data
        return {
            'chunk_id': chunk_id,
            'texts': [doc.page_content for doc in documents],
            'metadatas': [doc.metadata for doc in documents],
            'start_page': start_page
        }
        
    except Exception as e:
        logging.error(f"Error processing PDF chunk {chunk_id}: {str(e)}")
        # Clean up temp file if it exists
        temp_pdf_path = f"{pdf_path}_chunk_{chunk_id}.pdf"
        if os.path.exists(temp_pdf_path):
            os.remove(temp_pdf_path)
        raise

async def process_pdf_multiprocessing(filepath: str, filename: str, max_workers: int = None):
    """
    Process PDF using multiprocessing for maximum speed.
    """
    if max_workers is None:
        max_workers = min(mp.cpu_count(), 4)
    
    try:
        # Get PDF page count
        doc = fitz.open(filepath)
        total_pages = doc.page_count
        doc.close()
        
        if total_pages == 0:
            raise ValueError("PDF has no pages")
        
        # Calculate chunk size
        pages_per_chunk = max(1, total_pages // max_workers)
        
        # Prepare chunks for parallel processing
        chunks = []
        for i in range(0, total_pages, pages_per_chunk):
            chunks.append((
                pages_per_chunk,  # chunk_pages
                filepath,         # pdf_path
                i,               # start_page
                len(chunks)      # chunk_id
            ))
        
        logging.info(f"Processing {total_pages} pages in {len(chunks)} chunks with {max_workers} workers")
        
        # Process chunks in parallel using ProcessPoolExecutor
        loop = await asyncio.get_event_loop()
        
        # Use run_in_executor to run the blocking ProcessPoolExecutor operations
        with ProcessPoolExecutor(max_workers=max_workers) as executor:
            tasks = []
            for chunk in chunks:
                task = loop.run_in_executor(executor, process_pdf_chunk, chunk)
                tasks.append(task)
            
            # Wait for all chunks to complete
            chunk_results = await asyncio.gather(*tasks)
        
        # Combine all results
        all_texts = []
        all_metadatas = []
        
        for result in sorted(chunk_results, key=lambda x: x['chunk_id']):
            all_texts.extend(result['texts'])
            all_metadatas.extend(result['metadatas'])
        
        logging.info(f"Successfully processed {len(all_texts)} text chunks from {filename}")
        
        return all_texts, all_metadatas
        
    except Exception as e:
        logging.error(f"Error in multiprocessing PDF: {str(e)}")
        raise

async def query_pdf_context(query: str) -> dict:
    """
    Extract relevant context from PDF documents.
    """
    try:
        # Use existing Qdrant search
        result = qdrant_index.query_and_generate_response(query=query)
        
        if isinstance(result, tuple) and len(result) >= 2:
            context, sources = result[0], result[1]
        else:
            context = str(result)
            sources = []
        
        # Extract source information
        source_info = []
        if hasattr(sources, '__iter__'):
            for source in sources:
                if hasattr(source, 'page_content') and hasattr(source, 'metadata'):
                    source_info.append({
                        "content": source.page_content[:200] + "...",
                        "metadata": source.metadata
                    })
        
        return {
            "context": context,
            "sources": source_info
        }
        
    except Exception as e:
        logging.error(f"Error querying PDF context: {str(e)}")
        return {"context": "", "sources": []}

async def generate_comprehensive_response(query: str, pdf_context: str = "", max_tokens: int = 1000) -> str:
    """
    Generate either a comprehensive response or a set of questions with answers from the PDF context.
    Detects if the query is about question generation and adapts the prompt accordingly.
    """
    try:
        from openai import AsyncOpenAI

        client = AsyncOpenAI(api_key=settings.openai_api_key)

        is_question_generation = "generate" in query.lower() and "question" in query.lower()

        if is_question_generation:
            user_prompt = f"""
You are given the following PDF content. Based on this, carry out the task described in the query below.

**Query Instruction:**
{query}

**PDF Content:**
\"\"\"
{pdf_context}
\"\"\"

Please generate the number of questions specified in the query. Each question must be:
- Well-structured and clearly phrased
- Aligned with the academic content provided
- Diverse in format (e.g., MCQ, short answer, descriptive)
- Varied in difficulty level (definition, analysis, application)
- Accompanied by a detailed and accurate answer

Ensure your output adheres strictly to the required format and reflects only the information from the provided document.
"""

        else:
            system_prompt = """You are an expert AI assistant that provides comprehensive, detailed, and accurate responses.
Your goal is to give complete information about the query, drawing from your knowledge base and any provided context.

Guidelines:
1. Provide detailed, well-structured responses
2. If PDF context is provided, integrate it seamlessly
3. Use clear headings and bullet points for complex topics
4. Cite context where appropriate
"""

            user_prompt = f"""Query: {query}

{f"PDF Context:\n{pdf_context}" if pdf_context else "No PDF context available - please provide comprehensive information."}

Please provide a thorough response that fully addresses the query.
"""

        # Call OpenAI
        response = await client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": system_prompt.strip()},
                {"role": "user", "content": user_prompt.strip()}
            ],
            max_tokens=max_tokens,
            temperature=0.5
        )

        return response.choices[0].message.content

    except Exception as e:
        logging.error(f"Error generating comprehensive response: {str(e)}")
        if pdf_context:
            return f"[Fallback] Based on the PDF context:\n\n{pdf_context}\n\nQuery: {query}\n\n[LLM failed to generate a structured answer.]"
        else:
            return f"[Fallback] Query: '{query}'\nNo PDF context was provided. Unable to generate response due to system error."

async def combine_pdf_and_ai_insights(query: str, pdf_context: str, ai_response: str) -> dict:
    """
    Combine insights from PDF context and AI response to provide structured information.
    """
    try:
        return {
            "query_analysis": {
                "main_topic": extract_main_topic(query),
                "question_type": classify_question_type(query)
            },
            "pdf_insights": {
                "context_available": bool(pdf_context),
                "context_length": len(pdf_context) if pdf_context else 0,
                "key_points": extract_key_points_from_context(pdf_context) if pdf_context else []
            },
            "ai_insights": {
                "response_length": len(ai_response),
                "main_sections": extract_sections_from_response(ai_response)
            },
            "combined_confidence": "high" if pdf_context else "medium"
        }
    except Exception as e:
        logging.error(f"Error combining insights: {str(e)}")
        return {"error": "Could not combine insights", "details": str(e)}

def extract_main_topic(query: str) -> str:
    """Extract the main topic from a query using simple keyword analysis."""
    # Simple implementation - can be enhanced with NLP
    words = query.lower().split()
    # Remove common question words
    stop_words = {'what', 'how', 'why', 'when', 'where', 'who', 'is', 'are', 'the', 'a', 'an'}
    keywords = [word for word in words if word not in stop_words and len(word) > 2]
    return ' '.join(keywords[:3])  # Return first 3 keywords

def extract_sections_from_response(response: str) -> list:
    """Extract main sections from AI response."""
    sections = []
    lines = response.split('\n')
    
    for line in lines:
        line = line.strip()
        # Look for headers (lines that might be section titles)
        if line and (line.isupper() or line.startswith('#') or line.endswith(':')):
            sections.append(line)
    
    return sections  # Return first 5 sections


def format_for_frontend(response_text: str) -> dict:
    """
    Format response text for better frontend display.
    """
    try:
        # Clean the response
        cleaned_text = simple_clean_text(response_text)
        
        # Split into paragraphs
        paragraphs = [p.strip() for p in cleaned_text.split('\n\n') if p.strip()]
        
        # Extract main answer (first substantial paragraph)
        main_answer = paragraphs[0] if paragraphs else cleaned_text
        
        # Look for bullet points or numbered lists
        bullet_points = []
        numbered_points = []
        
        for para in paragraphs:
            if para.startswith(('•', '-', '*')):
                bullet_points.append(para)
            elif re.match(r'^\d+\.', para):
                numbered_points.append(para)
        
        return {
            "main_answer": main_answer,
            "full_response": cleaned_text,
            "paragraphs": paragraphs,
            "bullet_points": bullet_points,
            "numbered_points": numbered_points,
            "word_count": len(cleaned_text.split()),
            "has_structure": bool(bullet_points or numbered_points)
        }
        
    except Exception as e:
        logging.error(f"Error formatting response: {str(e)}")
        return {
            "main_answer": response_text,
            "full_response": response_text,
            "paragraphs": [response_text],
            "bullet_points": [],
            "numbered_points": [],
            "word_count": len(response_text.split()),
            "has_structure": False
        }

def simple_clean_text(text: str) -> str:
    """
    Simple text cleaning function.
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = ' '.join(text.split())
    
    # Remove or replace problematic characters
    text = text.replace('\x00', '')  # Remove null characters
    text = text.replace('\r\n', '\n')  # Normalize line endings
    text = text.replace('\r', '\n')
    
    # Remove excessive newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()

def cleanup_old_files():
    """
    Clean up old temporary files and expired sessions.
    """
    try:
        documents_dir = os.path.join("app", "documents")
        if os.path.exists(documents_dir):
            # Clean up any temporary chunk files that might have been left behind
            for filename in os.listdir(documents_dir):
                if "_chunk_" in filename and filename.endswith(".pdf"):
                    file_path = os.path.join(documents_dir, filename)
                    try:
                        # Remove files older than 1 hour
                        if time.time() - os.path.getmtime(file_path) > 3600:
                            os.remove(file_path)
                            logging.info(f"Cleaned up old temp file: {filename}")
                    except Exception as e:
                        logging.warning(f"Could not clean up temp file {filename}: {e}")
        
        # Clean up expired sessions
        cleanup_expired_sessions()
        
    except Exception as e:
        logging.error(f"Error in cleanup_old_files: {str(e)}")

# Public endpoints (no authentication required)
@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.utcnow().isoformat(),
        "version": "1.0.0"
    }
    
def clean_text_for_docx(text: str) -> str:
    """
    Remove null bytes and control characters that are not XML-compatible.
    """
    if not text:
        return ""
    # Remove null bytes and control characters (except newline/tab)
    return ''.join(c for c in text if c.isprintable() or c in '\n\t').replace('\x00', '').strip()

    
# Updated download-file route
@app.get("/download-conversations")
async def download_conversations(
    request: Request,
    user_id: int = None,
    limit: int = 100
):
    """Download conversation history as Word document"""
    try:
        # Check for authentication cookie
        session_token = request.cookies.get(COOKIE_NAME)
        if not session_token:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Authentication required. Please login."
            )
        
        # Validate session token
        current_user = get_user_by_session(session_token)
        if not current_user:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Invalid or expired session. Please login again."
            )
        
        # Get conversation history
        conversations = get_conversation_history(user_id, limit)
        
        if not conversations:
            raise HTTPException(status_code=404, detail="No conversations found")
        
        # Determine user name
        user_name = current_user.get('username', 'User')
        
        # Create Word document
        doc = create_conversation_document(conversations, user_name)
        
        # Save to memory buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"conversation_history_{timestamp}.docx"
        
        # Create streaming response
        def generate():
            yield buffer.read()
        
        return StreamingResponse(
            io.BytesIO(buffer.read()),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Error downloading conversations: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")
    
    
@app.get("/public/info")
async def public_info():
    """Public information about the API"""
    return {
        "name": "DrQA Backend API",
        "description": "PDF Question-Answering system with user authentication",
        "features": [
            "User registration and authentication",
            "PDF upload and processing",
            "Intelligent question answering",
            "Document management"
        ],
        "authentication_required": True
    }

# Error handlers
@app.exception_handler(404)
async def not_found_handler(request: Request, exc):
    return JSONResponse(
        status_code=404,
        content={"message": "Resource not found", "path": request.url.path}
    )

@app.exception_handler(500)
async def internal_error_handler(request: Request, exc):
    logging.error(f"Internal server error: {str(exc)}")
    return JSONResponse(
        status_code=500,
        content={"message": "Internal server error", "detail": "Please try again later"}
    )

# Background task for periodic cleanup
async def periodic_cleanup():
    """Periodic cleanup task"""
    while True:
        try:
            cleanup_old_files()
            await asyncio.sleep(3600)  # Run every hour
        except Exception as e:
            logging.error(f"Error in periodic cleanup: {str(e)}")
            await asyncio.sleep(300)  # Wait 5 minutes before retrying

def extract_key_points_from_context(context: str) -> list:
    """Extract key points from PDF context."""
    if not context:
        return []
    
    # Simple implementation - split by sentences and take important ones
    sentences = context.split('.')
    key_points = []
    
    for sentence in sentences[:5]:  # Take first 5 sentences
        sentence = sentence.strip()
        if len(sentence) > 20:  # Only meaningful sentences
            key_points.append(sentence + '.')
    
    return key_points


# Run the server
if __name__ == "__main__":
    # Start periodic cleanup task
    # asyncio.create_task(periodic_cleanup())
    
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        workers=1  # Single worker to avoid multiprocessing issues
    )
    
    